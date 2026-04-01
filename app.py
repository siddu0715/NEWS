import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
from PIL import Image
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False
import os
import yaml

# Set page config for a premium feel
st.set_page_config(
    page_title="News Intelligence System",
    page_icon="📰",
    layout="wide",
)

# Custom CSS for aesthetics
st.markdown("""
<style>
    .main { background-color: #0e1117; }
    .stButton>button { width: 100%; height: 3em; background-color: #4CAF50; color: white; border-radius: 8px; }
    .sentiment-good { color: #4CAF50; font-weight: bold; font-size: 1.2em; }
    .sentiment-average { color: #FFC107; font-weight: bold; font-size: 1.2em; }
    .sentiment-bad { color: #F44336; font-weight: bold; font-size: 1.2em; }
    .sidebar .stTextInput>div>div>input { border-color: #4CAF50; }
</style>
""", unsafe_allow_html=True)

# Define functions FIRST
def initialize_gemini(api_key):
    genai.configure(api_key=api_key)
    return genai.GenerativeModel('gemini-2.5-flash')

def analyze_content(model, parts):
    prompt = """
    Analyze the provided content (image/document pages). It may be in Telugu or English.
    Perform the following:
    1. Extract the Main Headline.
    2. Extract a comprehensive Summary in 5-8 descriptive bullet points.
    3. Sentiment Analysis: Classify as 'Good', 'Average', or 'Bad' based on the tone towards the government.
    4. Sentiment Reasoning: Provide a detailed explanation of the logic.

    Return the result strictly in this YAML format:
    Headline: "[Text]"
    Summary:
      - "[Point 1]"
      - "[Point 2]"
    Sentiment: "[Good/Average/Bad]"
    Reasoning: "[Detailed Description]"
    """
    response = model.generate_content([prompt] + parts)
    return response.text

def create_docx(data):
    doc = Document()
    doc.add_heading('Automated News Intelligence Report', 0)
    
    doc.add_heading('Headline', level=1)
    doc.add_paragraph(data.get('Headline', 'N/A'))
    
    doc.add_heading('Key Points', level=1)
    for point in data.get('Summary', []):
        doc.add_paragraph(point, style='List Bullet')
        
    doc.add_heading('Sentiment Analysis', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Tone: ").bold = True
    p.add_run(data.get('Sentiment', 'N/A'))
    
    doc.add_paragraph(f"Reasoning: {data.get('Reasoning', 'N/A')}")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def extract_pdf_pages(uploaded_file):
    if not FITZ_AVAILABLE:
        st.error("PDF processing is currently unavailable (missing 'pymupdf' module). Please contact the developer.")
        return []
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    images = []
    for page in doc:
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Better resolution
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    return images

def perform_analysis(parts, api_key):
    if not api_key:
        st.warning("Please enter your API key in the sidebar.")
        return

    try:
        with st.spinner("Decoding content and analyzing..."):
            model = initialize_gemini(api_key)
            result_text = analyze_content(model, parts)
            
            clean_text = result_text.replace("```yaml", "").replace("```", "").strip()
            data = yaml.safe_load(clean_text)

            if data:
                st.balloons()
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    st.success("### Results")
                    st.markdown(f"#### 📰 Headline: {data.get('Headline')}")
                    st.markdown("---")
                    st.write("**Key Highlights:**")
                    for p in data.get('Summary', []):
                        st.write(f"🔹 {p}")
                
                with col2:
                    sentiment = data.get('Sentiment', 'Average')
                    sentiment_label = sentiment or "Average"
                    color_class = f"sentiment-{sentiment_label.lower().strip()}"
                    st.markdown(f"### Tone Score: <span class='{color_class}'>{sentiment_label}</span>", unsafe_allow_html=True)
                    st.info(f"**AI Reasoning:**\n\n{data.get('Reasoning')}")
                
                # Word Doc download
                docx_file = create_docx(data)
                st.download_button(
                    label="📥 Export to Word Document (.docx)",
                    data=docx_file,
                    file_name=f"Report_{data.get('Headline', 'News')[:15]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    except Exception as e:
        err = str(e)
        if '429' in err or 'quota' in err.lower():
            st.warning("⏳ **API Quota Exceeded.** You've hit the free-tier rate limit.")
            st.info("**How to fix:** Wait ~1 minute, then click the button again. Alternatively, enable billing at [aistudio.google.com](https://aistudio.google.com/) for higher limits.")
        else:
            st.error(f"Analysis Failed: {e}")

# Main App UI
st.title("📰 News Intelligence System")
st.subheader("Transforming Print Media into Actionable Intelligence")

with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("Enter Gemini API Key", value="AIzaSyC_iQOPG_io86xYQkwJk41nPLoHXSiBzfs", type="password")
    st.info("No Key? Get it for free: [aistudio.google.com](https://aistudio.google.com/)")
    st.divider()
    st.markdown("### Supported Formats")
    st.write("✅ Images: JPG, PNG, JPEG")
    st.write("✅ Documents: PDF (Multi-page support)")

# Main Interface Tabs
tab1, tab2 = st.tabs(["🖼️ Image Analysis", "📄 PDF/Doc Analysis"])

with tab1:
    img_file = st.file_uploader("Upload Image", type=["jpg", "jpeg", "png"], key="img_up")
    if img_file:
        st.image(img_file, caption="Selected Image", use_container_width=True)
        if st.button("🚀 Analyze Image", key="btn_img"):
            process_data = [Image.open(img_file)]
            perform_analysis(process_data, api_key)

with tab2:
    if not FITZ_AVAILABLE:
        st.warning("⚠️ PDF processing requires the 'pymupdf' library. Please check your environment.")
    pdf_file = st.file_uploader("Upload PDF Document", type=["pdf"], key="pdf_up")
    if pdf_file and FITZ_AVAILABLE:
        st.success(f"File uploaded: {pdf_file.name}")
        if st.button("📑 Process Document", key="btn_pdf"):
            with st.spinner("Extracting pages and analyzing multi-page content..."):
                images = extract_pdf_pages(pdf_file)
                if images:
                    perform_analysis(images, api_key)

if not api_key:
    st.info("👈 Please start by entering your Google Gemini API Key in the settings.")
